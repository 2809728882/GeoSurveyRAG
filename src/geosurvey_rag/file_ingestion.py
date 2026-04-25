from __future__ import annotations

import csv
import io
import json
from pathlib import Path
from zipfile import BadZipFile

from geosurvey_rag.knowledge_sources import content_sha1, now_iso, slugify
from geosurvey_rag.settings import settings


SUPPORTED_UPLOAD_EXTENSIONS = {
    ".md",
    ".txt",
    ".csv",
    ".tsv",
    ".json",
    ".pdf",
    ".docx",
    ".xlsx",
}


def decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "gbk"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def parse_plain_text(data: bytes) -> str:
    return decode_text(data)


def parse_csv(data: bytes, delimiter: str = ",") -> str:
    text = decode_text(data)
    rows = list(csv.reader(io.StringIO(text), delimiter=delimiter))
    lines = ["| " + " | ".join(cell.strip() for cell in row) + " |" for row in rows[:500]]
    return "\n".join(lines)


def parse_json(data: bytes) -> str:
    obj = json.loads(decode_text(data))
    return json.dumps(obj, ensure_ascii=False, indent=2)


def parse_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("PDF support requires pypdf.") from exc

    reader = PdfReader(io.BytesIO(data))
    pages = []
    for index, page in enumerate(reader.pages, start=1):
        pages.append(f"## Page {index}\n\n{page.extract_text() or ''}")
    return "\n\n".join(pages)


def parse_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("DOCX support requires python-docx.") from exc

    document = Document(io.BytesIO(data))
    parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def parse_xlsx(data: bytes) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise RuntimeError("XLSX support requires openpyxl.") from exc

    workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts = []
    for sheet in workbook.worksheets:
        parts.append(f"# Sheet: {sheet.title}")
        for row in sheet.iter_rows(max_row=500, values_only=True):
            values = ["" if value is None else str(value) for value in row]
            if any(values):
                parts.append("| " + " | ".join(values) + " |")
    return "\n".join(parts)


def extract_upload_text(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix in {".md", ".txt"}:
        return parse_plain_text(data)
    if suffix == ".csv":
        return parse_csv(data)
    if suffix == ".tsv":
        return parse_csv(data, delimiter="\t")
    if suffix == ".json":
        return parse_json(data)
    if suffix == ".pdf":
        return parse_pdf(data)
    if suffix == ".docx":
        return parse_docx(data)
    if suffix == ".xlsx":
        return parse_xlsx(data)
    raise ValueError(f"Unsupported upload file type: {suffix}")


def save_uploaded_knowledge(
    filename: str,
    data: bytes,
    category: str = "uploads",
    knowledge_dir: Path = settings.knowledge_dir,
) -> dict:
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_UPLOAD_EXTENSIONS:
        raise ValueError(f"Unsupported upload file type: {suffix}")

    text = extract_upload_text(filename, data).strip()
    if not text:
        raise ValueError(f"No extractable text found in {filename}")

    upload_id = content_sha1(filename + content_sha1(text))[:12]
    safe_stem = slugify(Path(filename).stem)
    raw_dir = Path("data/uploads") / slugify(category)
    raw_dir.mkdir(parents=True, exist_ok=True)
    raw_path = raw_dir / f"{safe_stem}-{upload_id}{suffix}"
    raw_path.write_bytes(data)

    target_dir = knowledge_dir / "uploads" / slugify(category)
    target_dir.mkdir(parents=True, exist_ok=True)
    target = target_dir / f"{safe_stem}-{upload_id}.md"
    body = [
        "---",
        "source_type: upload",
        f"filename: {filename}",
        f"category: {category}",
        f"raw_path: {raw_path}",
        f"uploaded_at: {now_iso()}",
        f"content_sha1: {content_sha1(text)}",
        "---",
        "",
        f"# Uploaded: {filename}",
        "",
        text,
        "",
    ]
    target.write_text("\n".join(body), encoding="utf-8")
    return {
        "filename": filename,
        "ok": True,
        "knowledge_path": str(target),
        "raw_path": str(raw_path),
        "chars": len(text),
        "content_sha1": content_sha1(text),
    }
